import pandas as pd
import docx
from docx.shared import RGBColor, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import re
import os
import glob

# --- Configuration ---
CSV_FILE = 'HKDSE_17to25_HKDSE_verified_data.csv'

# --- Chosen Pale Color Palette (From Attachment) ---
COLOR_GREEN = {
    "text": RGBColor(0, 0, 0),      # Black text
    "fill": "E2EFDA",               # Green fill from attachment
    "border": "A9D08E"              # Green border from attachment
}
COLOR_ORANGE = {
    "text": RGBColor(0, 0, 0),      # Black text
    "fill": "FFF2CC",               # Yellow fill from attachment
    "border": "FFD966"              # Yellow border from attachment
}
COLOR_RED = {
    "text": RGBColor(0, 0, 0),      # Black text
    "fill": "FCE4D6",               # Red/Pink fill from attachment
    "border": "F4B084"              # Red/Pink border from attachment
}

# --- Score Thresholds ---
THRESHOLD_HIGH = 60 # 60%+ Green
THRESHOLD_MID = 40  # 40-59% Orange, <40% Red

def get_style_for_score(score_str):
    try:
        score = int(str(score_str).replace('%', '').strip())
        if score >= THRESHOLD_HIGH:
            return COLOR_GREEN
        elif score >= THRESHOLD_MID:
            return COLOR_ORANGE
        else:
            return COLOR_RED
    except ValueError:
        return COLOR_GREEN

def apply_boxed_style(run, style_dict):
    """Injects exact XML to apply border, shading, and text color to a run."""
    run.font.color.rgb = style_dict["text"]
    rPr = run._r.get_or_add_rPr()
    
    shd_xml = f'<w:shd {nsdecls("w")} w:fill="{style_dict["fill"]}" w:val="clear"/>'
    rPr.append(parse_xml(shd_xml))
    
    bdr_val = f'<w:rBdr {nsdecls("w")} w:val="single" w:sz="4" w:space="0" w:color="{style_dict["border"]}"/>'
    rPr.append(parse_xml(bdr_val))

def process_document(doc_path, target_year, target_paper, df):
    """Processes a single Word document."""
    print(f"\n--- Processing: {doc_path} (Year: {target_year}, Paper: {target_paper}) ---")
    
    # Filter CSV for this specific document's year and paper
    filtered_df = df[(df['Year'] == target_year) & (df['Paper'] == target_paper)]
    
    grouped_stats = {}
    for k, v in zip(filtered_df['Question No.'], filtered_df['HK % score']):
        clean_key = str(k).replace(" ", "").strip()
        score_str = str(v).strip()
        
        match = re.match(r'^(\d+)(.*)$', clean_key)
        if match:
            q_num = match.group(1)
            sub_parts = match.group(2)
            if q_num not in grouped_stats:
                grouped_stats[q_num] = []
            grouped_stats[q_num].append((sub_parts, score_str))

    if not grouped_stats:
        print(f"  -> Skipping: No data found in CSV for Year {target_year}, Paper {target_paper}.")
        return

    doc = docx.Document(doc_path)
    
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
            
        q_match = re.search(r'^\W*(?:Q|Question)?\s*(\d+)\s*[\.\)]?', text, re.IGNORECASE)
        if q_match:
            q_num = q_match.group(1)
            
            if q_num in grouped_stats:
                stats_p = p.insert_paragraph_before()
                
                prefix_run = stats_p.add_run(f"Q{q_num}   ")
                prefix_run.bold = True
                prefix_run.font.size = Pt(11)
                prefix_run.font.color.rgb = RGBColor(0, 0, 0)
                
                stats_list = grouped_stats[q_num]
                for sub_parts, score_str in stats_list:
                    exact_style = get_style_for_score(score_str)
                    stat_display_text = f" {sub_parts}: {score_str} "
                    
                    stat_run = stats_p.add_run(stat_display_text)
                    stat_run.bold = True
                    stat_run.font.size = Pt(11)
                    apply_boxed_style(stat_run, exact_style)
                    
                    stats_p.add_run("   ")
                
                stats_p.add_run("\n")
                del grouped_stats[q_num]

    # Save output file
    output_filename = doc_path.replace('.docx', '_Annotated.docx')
    doc.save(output_filename)
    print(f"  -> Success! Saved as: {output_filename}")
    
    if grouped_stats:
        print(f"  -> Warning: The following questions were in the CSV but not found in the document:")
        for q_num, parts in grouped_stats.items():
            print(f"     - Question {q_num}")

def main():
    print("Loading data from CSV...")
    try:
        df = pd.read_csv(CSV_FILE)
    except FileNotFoundError:
        print(f"Error: Could not find '{CSV_FILE}'. Please make sure it's in the same folder.")
        input("\nPress Enter to exit...")
        return

    # Find all .docx files in the current folder
    # Ignore temporary files (~$) and already annotated files to prevent double-processing
    docx_files = [f for f in glob.glob("*.docx") if not f.endswith("_Annotated.docx") and not f.startswith("~$")]
    
    if not docx_files:
        print("No Word documents found in this folder.")
        input("\nPress Enter to exit...")
        return

    # Process each file found
    for file in docx_files:
        # Auto-detect year and paper from filename (e.g., "HKDSE_2025_Paper 1B.docx" -> Year: 2025, Paper: 1B)
        match = re.search(r'HKDSE_(\d{4})_Paper\s*([A-Za-z0-9]+)', file, re.IGNORECASE)
        
        if match:
            target_year = int(match.group(1))
            target_paper = match.group(2).upper()
            
            process_document(file, target_year, target_paper, df)
        else:
            print(f"\n--- Skipping: {file} ---")
            print("  -> Could not detect Year and Paper from the filename. Please name it like 'HKDSE_2025_Paper 1B.docx'")

    print("\n========================================")
    print("Batch processing complete!")
    print("========================================")
    input("\nPress Enter to exit...")

if __name__ == "__main__":
    main()