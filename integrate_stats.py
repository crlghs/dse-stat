import pandas as pd
import docx
from docx.shared import RGBColor, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import re

# --- Configuration ---
CSV_FILE = 'HKDSE_17to25_HKDSE_verified_data.csv'
DOC_FILE = 'HKDSE_2022_Paper 1B.docx'
OUTPUT_FILE = 'HKDSE_2022_Paper 1B_Annotated.docx'
TARGET_YEAR = 2023  
TARGET_PAPER = '1B'

# --- Chosen Pale Color Palette (From Attachment) ---
COLOR_GREEN = {
    "text": RGBColor(0, 0, 0),      # Black text
    "fill": "E2EFDA",               # Green fill from attachment
    "border": "A9D08E"              # Green border from attachment
}
COLOR_ORANGE = {
    "text": RGBColor(0, 0, 0),      # Black text
    "fill": "FFF2CC",               # Yellow fill from attachment
    "border": "FFD966"              # Yellow border from attachment
}
COLOR_RED = {
    "text": RGBColor(0, 0, 0),      # Black text
    "fill": "FCE4D6",               # Red/Pink fill from attachment
    "border": "F4B084"              # Red/Pink border from attachment
}

# --- Score Thresholds (Adjustable) ---
THRESHOLD_HIGH = 60 # 60%+ Green
THRESHOLD_MID = 40  # 40-59% Orange, <40% Red

def get_style_for_score(score_str):
    try:
        score = int(str(score_str).replace('%', '').strip())
        if score >= THRESHOLD_HIGH:
            return COLOR_GREEN
        elif score >= THRESHOLD_MID:
            return COLOR_ORANGE
        else:
            return COLOR_RED
    except ValueError:
        return COLOR_GREEN # Fallback color

def apply_boxed_style(run, style_dict):
    """Injects exact XML to apply border, shading, and text color to a run."""
    run.font.color.rgb = style_dict["text"]
    rPr = run._r.get_or_add_rPr()
    
    # Fill color
    shd_xml = f'<w:shd {nsdecls("w")} w:fill="{style_dict["fill"]}" w:val="clear"/>'
    rPr.append(parse_xml(shd_xml))
    
    # Border color (sz="4" gives a clean, thin border line)
    bdr_val = f'<w:rBdr {nsdecls("w")} w:val="single" w:sz="4" w:space="0" w:color="{style_dict["border"]}"/>'
    rPr.append(parse_xml(bdr_val))

def main():
    print("Loading data from CSV...")
    df = pd.read_csv(CSV_FILE)
    filtered_df = df[(df['Year'] == TARGET_YEAR) & (df['Paper'] == TARGET_PAPER)]
    
    grouped_stats = {}
    for k, v in zip(filtered_df['Question No.'], filtered_df['HK % score']):
        clean_key = str(k).replace(" ", "").strip()
        score_str = str(v).strip()
        
        # Matches main question number and the sub-parts
        match = re.match(r'^(\d+)(.*)$', clean_key)
        if match:
            q_num = match.group(1)
            sub_parts = match.group(2)
            
            if q_num not in grouped_stats:
                grouped_stats[q_num] = []
            grouped_stats[q_num].append((sub_parts, score_str))

    if not grouped_stats:
        print(f"No data found for Year {TARGET_YEAR}, Paper {TARGET_PAPER}.")
        input("\nPress Enter to exit...")
        return

    print("Opening Word document...")
    doc = docx.Document(DOC_FILE)
    
    print("Processing paragraphs...")
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
            
        q_match = re.search(r'^\W*(?:Q|Question)?\s*(\d+)\s*[\.\)]?', text, re.IGNORECASE)
        if q_match:
            q_num = q_match.group(1)
            
            if q_num in grouped_stats:
                # 1. Create the new paragraph directly ABOVE the question
                stats_p = p.insert_paragraph_before()
                
                # 2. Add the main prefix as plain bold text: "Q1 "
                prefix_run = stats_p.add_run(f"Q{q_num}   ")
                prefix_run.bold = True
                prefix_run.font.size = Pt(11)
                prefix_run.font.color.rgb = RGBColor(0, 0, 0) # Black text
                
                # 3. Loop through sub-parts for this question
                stats_list = grouped_stats[q_num]
                for sub_parts, score_str in stats_list:
                    
                    exact_style = get_style_for_score(score_str)
                    
                    # Construct text WITHOUT the main question number: e.g. " (a): 68% "
                    stat_display_text = f" {sub_parts}: {score_str} "
                    
                    stat_run = stats_p.add_run(stat_display_text)
                    stat_run.bold = True
                    stat_run.font.size = Pt(11)
                    apply_boxed_style(stat_run, exact_style)
                    
                    # Add standard space between boxes
                    stats_p.add_run("   ")
                
                stats_p.add_run("\n")
                del grouped_stats[q_num]

    doc.save(OUTPUT_FILE)
    print(f"\nSuccess! Annotated document saved as: {OUTPUT_FILE}")
    
    if grouped_stats:
        print("\nNote: The following main questions were in the CSV but could not be matched:")
        for q_num, parts in grouped_stats.items():
            print(f"- Question {q_num} (has {len(parts)} subparts in CSV)")

    input("\nPress Enter to exit...")

if __name__ == "__main__":
    main()